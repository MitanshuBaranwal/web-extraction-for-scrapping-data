from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import requests
from bs4 import BeautifulSoup
from docx import Document
import html2text
import re


def apply_formatting(paragraph, line):
    # Apply appropriate format of bold or italic to the paragraph on the basis of syntax.
    if '**' in line:
        # Split line based on bold syntax '**'
        parts = line.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Add non-bold part to the paragraph
                paragraph.add_run(part)
            else:
                # Add bold part to the paragraph
                paragraph.add_run(part).bold = True

    elif '_' in line:
        # Split the line based on italic syntax '_'
        parts = line.split('_')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Add non-italic part to the paragraph
                paragraph.add_run(part)
            else:
                # Add italic part to the paragraph
                paragraph.add_run(part).italic = True
    else:
        paragraph.add_run(line)


@csrf_exempt
def extract_text(request):
    # View function to extract a given URL process text and save as a Word document.

    if request.method == 'POST':
        # Retrieve the URL from the POST data
        url = request.POST.get('url')

        # Fetch HTML content from the provided URL
        response = requests.get(url)

        if response.status_code == 200:
            # Parse the HTML content using BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')

            # Convert HTML to Markdown with basic styling preserved
            h = html2text.HTML2Text()
            h.body_width = 0
            h.inline_links = False  # Preserve links as-is
            markdown_text = h.handle(str(soup))

            # Remove non-needed URLs from the content
            url_regex = r'\[.*?]\(.*?\)'
            markdown_text_no_urls = re.sub(url_regex,' ', markdown_text)

            # Create a new Word document
            doc = Document()
            for line in markdown_text_no_urls.splitlines():
                p = doc.add_paragraph()

                # Handle headers and text formatting
                if line.startswith('#'):
                    doc.add_heading(line.lstrip('# '), level=1)
                else:
                    apply_formatting(p, line)
            print(p)
            # Save the Word document
            file_path = 'extracted_text.docx'
            doc.save(file_path)

            # Return a success message
            return JsonResponse({'message': 'Text extracted and saved as extracted_text.docx'})
        else:
            # Return an error message if fetching the URL failed
            return JsonResponse({'error': 'Failed to fetch URL.'}, status=400)
    else:
        # Return an error message for unsupported HTTP methods
        return JsonResponse({'error': 'Only POST requests are supported.'}, status=405)
